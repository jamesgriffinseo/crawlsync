#!/usr/bin/env python3
"""
Sitemap extraction server for CrawlSync.

Install deps:  pip install flask flask-cors requests
Run:           python sitemap_server.py
Then open:     sitemap + IA v2.html
"""

import os
import re
import sys
import time
import gzip as _gzip
import threading
import xml.etree.ElementTree as ET
from urllib.parse import urlparse
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import requests
try:
    import cloudscraper as _cloudscraper
    _SCRAPER = _cloudscraper.create_scraper()
except Exception:
    _SCRAPER = None

try:
    from curl_cffi import requests as _cffi_requests
    _CFFI = True
except Exception:
    _CFFI = False

# ── Headless browser (Playwright) — persistent background instance ───────────
# Launched once on first use; shared across all requests via a threading lock.
# Falls back gracefully when Playwright / Chromium is not available.
try:
    from playwright.sync_api import sync_playwright as _sync_playwright
    _PW_AVAILABLE = True
except ImportError:
    _PW_AVAILABLE = False

_pw_lock     = threading.Lock()
_pw_instance = None   # playwright context manager
_pw_browser  = None   # persistent Browser object


def _get_pw_browser():
    """Return the shared Playwright Browser, launching it if needed."""
    global _pw_instance, _pw_browser
    if not _PW_AVAILABLE:
        return None
    with _pw_lock:
        if _pw_browser is None or not _pw_browser.is_connected():
            try:
                _pw_instance = _sync_playwright().start()
                _pw_browser  = _pw_instance.chromium.launch(
                    headless=True,
                    args=["--no-sandbox", "--disable-dev-shm-usage",
                          "--disable-blink-features=AutomationControlled"],
                )
            except Exception as e:
                print(f"[Playwright] Failed to launch: {e}")
                return None
        return _pw_browser


def _render_with_playwright(url, timeout_ms=20000):
    """Render a URL in headless Chromium and return the page HTML.

    Opens a fresh BrowserContext per call (separate cookies/cache) so
    requests don't bleed into each other.  Returns None on failure.
    """
    browser = _get_pw_browser()
    if not browser:
        return None
    ctx  = None
    page = None
    try:
        ctx  = browser.new_context(
            user_agent="Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                       "AppleWebKit/537.36 (KHTML, like Gecko) "
                       "Chrome/124.0.0.0 Safari/537.36",
            java_script_enabled=True,
            bypass_csp=True,
        )
        page = ctx.new_page()
        # Block heavy resources that don't affect content (fonts, media, analytics)
        def _block(route):
            if route.request.resource_type in ("image", "media", "font"):
                route.abort()
            else:
                route.continue_()
        page.route("**/*", _block)
        page.goto(url, wait_until="domcontentloaded", timeout=timeout_ms)
        # Wait for visible body text — up to 10 s
        try:
            page.wait_for_function(
                "() => document.body && document.body.innerText.trim().length > 100",
                timeout=10000,
            )
        except Exception:
            pass   # use whatever rendered so far
        html = page.content()
        return html
    except Exception as e:
        print(f"[Playwright] Render failed for {url}: {e}")
        return None
    finally:
        try:
            if page:  page.close()
            if ctx:   ctx.close()
        except Exception:
            pass


def resource_path(rel):
    # Search all candidate locations so the app works regardless of which
    # directory PyInstaller sets sys._MEIPASS to (MacOS vs Frameworks vs Resources).
    candidates = []
    if hasattr(sys, "_MEIPASS"):
        meipass = sys._MEIPASS
        candidates += [
            meipass,
            os.path.normpath(os.path.join(meipass, "..", "Frameworks")),
            os.path.normpath(os.path.join(meipass, "..", "Resources")),
        ]
    candidates.append(os.path.dirname(os.path.abspath(__file__)))
    for base in candidates:
        path = os.path.join(base, rel)
        if os.path.exists(path):
            return path
    # Fallback: return the _MEIPASS path even if missing (preserves old behaviour)
    return os.path.join(candidates[0] if candidates else ".", rel)


app = Flask(__name__)
CORS(app)
app.config["MAX_CONTENT_LENGTH"] = 100 * 1024 * 1024   # 100 MB — allow large bulk-save payloads


@app.route("/debug-paths")
def debug_paths():
    import traceback
    html = "sitemap + IA v2.html"
    info = {
        "_MEIPASS": getattr(sys, "_MEIPASS", "NOT SET"),
        "__file__": __file__,
        "resource_path_result": resource_path(html),
        "file_exists": os.path.exists(resource_path(html)),
    }
    try:
        info["file_size"] = os.path.getsize(resource_path(html))
    except Exception as e:
        info["file_size_error"] = str(e)
    return jsonify(info)


@app.route("/")
def index():
    try:
        path = resource_path("sitemap + IA v2.html")
        return send_file(path)
    except Exception as e:
        import traceback
        return f"<pre>Error: {e}\n\nPath: {path}\nExists: {os.path.exists(path)}\n\n{traceback.format_exc()}</pre>", 500

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8",
    "Accept-Encoding": "gzip, deflate",
    "Accept-Language": "en-US,en;q=0.9",
    "Connection": "keep-alive",
    "Cache-Control": "no-cache",
    "Pragma": "no-cache",
    "Sec-Fetch-Dest": "document",
    "Sec-Fetch-Mode": "navigate",
    "Sec-Fetch-Site": "none",
    "Sec-Fetch-User": "?1",
    "Upgrade-Insecure-Requests": "1",
}

# Fallback — plain crawler identity, accepted by some WAFs/WordPress
HEADERS_CRAWLER = {
    "User-Agent": "CrawlSync/1.0 (Sitemap Crawler; +https://crawlsync.app)",
    "Accept": "application/xml,text/xml,*/*",
    "Accept-Encoding": "gzip, deflate",
}

# Googlebot — whitelisted by many Cloudflare configs
HEADERS_GOOGLEBOT = {
    "User-Agent": "Mozilla/5.0 (compatible; Googlebot/2.1; +http://www.google.com/bot.html)",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Encoding": "gzip, deflate",
}


def _decode_response(r):
    """Try multiple encodings to robustly decode a requests-like Response."""
    import zlib as _zlib
    raw = r.content
    # Manually decompress if auto-decompression didn't fire.
    # IMPORTANT: on failure return None — do NOT fall through with raw compressed
    # bytes, as latin-1 would silently decode them to garbage and lxml would then
    # throw zlib.error -3 ("incorrect header check") when it tries to parse it.
    if raw[:2] == b'\x1f\x8b':  # gzip magic
        try:
            raw = _gzip.decompress(raw)
        except Exception:
            return None   # unrecoverable — let caller handle missing content
    elif len(raw) >= 2 and raw[0] == 0x78 and raw[1] in (0x01, 0x5e, 0x9c, 0xda):  # zlib/deflate
        try:
            raw = _zlib.decompress(raw)
        except Exception:
            try:
                raw = _zlib.decompress(raw, -15)  # raw deflate (no header)
            except Exception:
                return None
    xml_enc = None
    try:
        header = raw[:300].decode("ascii", errors="ignore")
        m = re.search(r'encoding=["\']([^"\']+)["\']', header)
        if m:
            xml_enc = m.group(1)
    except Exception:
        pass
    apparent = getattr(r, "apparent_encoding", None)  # not present on curl_cffi responses
    for enc in filter(None, ["utf-8-sig", xml_enc, apparent, "utf-8", "latin-1"]):
        try:
            text = raw.decode(enc).strip()
            if text:
                return text
        except (UnicodeDecodeError, LookupError):
            continue
    return None


def _is_cloudflare_block(r):
    # Status-code gate: 403/429/503 are common CF block codes, but CF also
    # returns 200 with a JS-challenge page ("Enable JavaScript and cookies").
    if r.status_code not in (200, 403, 429, 503):
        return False
    body = r.text[:2000].lower()
    cf_markers = (
        "just a moment", "checking your browser", "_cf_chl",
        "challenge-platform", "cf-browser-verification", "cloudflare ray id",
        "enable javascript and cookies",
    )
    if not any(s in body for s in cf_markers):
        return False
    # For 200 responses, also require the page to be near-empty content-wise
    # so we don't false-positive on pages that legitimately mention Cloudflare.
    if r.status_code == 200:
        return len(body) < 20_000
    return True


def _try_cloudflare_bypass(url, timeout, log_lines):
    """Try progressively stronger bypass methods for Cloudflare-protected URLs."""
    def _log(msg):
        if log_lines is not None:
            log_lines.append(msg)

    bypass_timeout = min(timeout, 12)  # shorter timeout for bypass attempts

    # 1. curl_cffi — best Chrome TLS fingerprint mimic
    if _CFFI:
        try:
            r = _cffi_requests.get(url, impersonate="chrome120", timeout=bypass_timeout)
            text = _decode_response(r)
            if text and not _is_cloudflare_block(r):
                _log("  curl_cffi bypass succeeded")
                return text
        except Exception as e:
            _log(f"  curl_cffi failed: {e}")

    # 2. cloudscraper — JS challenge solver
    if _SCRAPER:
        try:
            r = _SCRAPER.get(url, timeout=bypass_timeout)
            text = _decode_response(r)
            if text and not _is_cloudflare_block(r):
                _log("  cloudscraper bypass succeeded")
                return text
        except Exception as e:
            _log(f"  cloudscraper failed: {e}")

    # 3. Googlebot UA — whitelisted in many Cloudflare configs
    try:
        r = requests.get(url, headers=HEADERS_GOOGLEBOT, timeout=bypass_timeout, allow_redirects=True)
        if r.ok and not _is_cloudflare_block(r):
            text = _decode_response(r)
            if text:
                _log("  Googlebot UA bypass succeeded")
                return text
    except Exception:
        pass

    _log("  All bypass methods failed — try Manual Setup to paste URLs directly")
    return None


def fetch(url, timeout=15, retries=1, log_lines=None):  # reduced defaults
    def _log(msg):
        if log_lines is not None:
            log_lines.append(msg)

    for headers in [HEADERS, HEADERS_CRAWLER]:
        for attempt in range(retries):
            try:
                r = requests.get(url, headers=headers, timeout=timeout, allow_redirects=True)
                if _is_cloudflare_block(r):
                    _log("  Cloudflare protection detected — trying bypass methods")
                    return _try_cloudflare_bypass(url, timeout, log_lines)
                if not r.ok:
                    _log(f"  HTTP {r.status_code} for {url}")
                    if attempt < retries - 1:
                        time.sleep(1)
                    continue
                text = _decode_response(r)
                if text:
                    return text
                _log(f"  Empty/undecodable response from {url}")
            except requests.exceptions.Timeout:
                _log(f"  Timeout (attempt {attempt + 1}) for {url}")
                if attempt < retries - 1:
                    time.sleep(1)
            except requests.exceptions.SSLError as e:
                _log(f"  SSL error: {e}")
                return None
            except requests.exceptions.ConnectionError as e:
                _log(f"  Connection error: {e}")
                return None
            except Exception as e:
                _log(f"  Error: {e}")
                return None

    return None


def looks_like_sitemap(url):
    """True if a URL is likely a sitemap rather than a regular page."""
    lower = url.split("?")[0].lower()  # strip query string for extension check
    full_lower = url.lower()
    return (
        "sitemap" in full_lower
        or lower.endswith(".xml")
        or lower.endswith(".xml.gz")
        or lower.endswith(".php")
        or lower.endswith(".asp")
        or lower.endswith(".aspx")
    )


def _is_xml_sitemap(text):
    """Return True only if the content looks like real XML/sitemap, not an HTML soft-404."""
    if not text:
        return False
    sniff = text.lstrip()[:600].lower()
    # Must start with XML declaration or contain sitemap-specific tags
    return (
        sniff.startswith("<?xml")
        or "<urlset" in sniff
        or "<sitemapindex" in sniff
        or ("<loc>" in sniff and "</loc>" in sniff)
    )


def _www_alt(origin):
    """Return the www↔non-www counterpart of an origin, or None if already tried."""
    p = urlparse(origin)
    if p.netloc.startswith("www."):
        return f"{p.scheme}://{p.netloc[4:]}"
    return f"{p.scheme}://www.{p.netloc}"


def _sitemap_www_alt(url):
    """Return the www↔non-www variant of a full sitemap URL."""
    p = urlparse(url)
    if p.netloc.startswith("www."):
        alt_netloc = p.netloc[4:]
    else:
        alt_netloc = "www." + p.netloc
    return url.replace(f"://{p.netloc}", f"://{alt_netloc}", 1)


def discover_sitemaps(raw_input, log_lines=None):
    """Return a list of reachable sitemap URLs for the given domain/URL."""
    def _log(msg):
        if log_lines is not None:
            log_lines.append(msg)

    url = raw_input.strip()
    if not url.startswith("http"):
        url = "https://" + url

    # If a direct sitemap URL was given, just use it
    if looks_like_sitemap(url) and "/" in url.replace("https://", "").replace("http://", ""):
        return [url]

    parsed = urlparse(url)
    origin     = f"{parsed.scheme}://{parsed.netloc}"
    alt_origin = _www_alt(origin)

    # ── 1. Collect sitemap URLs from robots.txt on BOTH www and non-www ──
    # Some sites declare sitemap only in one variant; check both.
    all_found = []
    seen_robots = set()
    for base in [origin, alt_origin]:
        robots_url = base + "/robots.txt"
        if robots_url in seen_robots:
            continue
        seen_robots.add(robots_url)
        robots = fetch(robots_url)
        if robots:
            hits = re.findall(r"Sitemap:\s*(https?://[^\s]+)", robots, re.IGNORECASE)
            for h in hits:
                if h not in all_found:
                    all_found.append(h)

    if all_found:
        _log(f"  Found {len(all_found)} sitemap declaration(s) in robots.txt")
        reachable = []
        for su in all_found:
            # Try declared URL first, then www↔non-www variant
            for candidate in dict.fromkeys([su, _sitemap_www_alt(su)]):
                content = fetch_sitemap(candidate)
                if _is_xml_sitemap(content):
                    if candidate != su:
                        _log(f"  {su} — soft-404; using alternate: {candidate}")
                    reachable.append(candidate)
                    break
            else:
                reason = "not reachable" if not fetch(su) else "soft-404 (HTML response)"
                _log(f"  {su} — {reason}, skipping")

        if reachable:
            return reachable
        _log(f"  All robots.txt sitemaps failed — trying common paths")

    # ── 2. Try common paths on both www and non-www origins ─────────────
    for base in dict.fromkeys([origin, alt_origin]):
        for path in FALLBACK_PATHS:
            candidate = base + path
            content = fetch_sitemap(candidate)
            if _is_xml_sitemap(content):
                _log(f"  Found sitemap via fallback: {candidate}")
                return [candidate]

    # Nothing found — return best guess so the caller can log a clear error
    return [origin + "/sitemap.xml"]


def parse_locs(text):
    """Extract all <loc> values from sitemap XML — handles namespaces, CDATA, and HTML entities."""
    # Handle CDATA: <loc><![CDATA[https://...]]></loc>
    text = re.sub(r"<!\[CDATA\[(.*?)\]\]>", r"\1", text, flags=re.DOTALL)
    locs = re.findall(r"<loc>\s*(https?://[^\s<]+)\s*</loc>", text, re.IGNORECASE)
    # Decode HTML entities (e.g. &amp; → &)
    locs = [l.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">").strip() for l in locs]
    return locs


def is_sitemap_index(text):
    """True if the document contains <sitemap> index entries."""
    return bool(re.search(r"<sitemap[\s>]", text, re.IGNORECASE))


def fetch_sitemap(url, timeout=15, log_lines=None):
    """Fetch a sitemap URL, preferring XML Accept headers to avoid XSL/HTML rendering."""
    def _log(msg):
        if log_lines is not None:
            log_lines.append(msg)

    # Try XML-preferring headers first (avoids Yoast XSL → HTML rendering)
    for headers in [HEADERS_CRAWLER, HEADERS]:
        try:
            r = requests.get(url, headers=headers, timeout=timeout, allow_redirects=True)
            if _is_cloudflare_block(r):
                _log("  Cloudflare protection detected — trying bypass methods")
                return _try_cloudflare_bypass(url, timeout, log_lines)
            if not r.ok:
                continue
            text = _decode_response(r)
            if text and _is_xml_sitemap(text):
                return text
            # Got a response but it's not XML — try next header set
            if text:
                _log(f"  Non-XML response with {headers.get('User-Agent','?')[:30]}… trying alternate headers")
        except requests.exceptions.SSLError as e:
            _log(f"  SSL error: {e}")
            return None
        except requests.exceptions.ConnectionError as e:
            _log(f"  Connection error: {e}")
            return None
        except Exception as e:
            _log(f"  Error: {e}")
            return None

    # Last resort: return whatever we can get
    return fetch(url, timeout=timeout, log_lines=log_lines)


def extract_urls(url, collected=None, visited=None, log_lines=None):
    if collected is None:
        collected = set()
    if visited is None:
        visited = set()
    if log_lines is None:
        log_lines = []
    if url in visited:
        return collected
    visited.add(url)

    log_lines.append(f"Scanning: {url}")
    text = fetch_sitemap(url, log_lines=log_lines)
    if not text:
        log_lines.append(f"Could not fetch: {url}")
        return collected

    locs = parse_locs(text)
    log_lines.append(f"  → {len(locs)} <loc> entries found")

    if is_sitemap_index(text):
        # It's an index — recurse into each child sitemap
        log_lines.append(f"  → sitemap index with {len(locs)} children")
        for loc in locs:
            if looks_like_sitemap(loc):
                extract_urls(loc, collected, visited, log_lines)
            else:
                # child loc doesn't look like a sitemap — try fetching it anyway
                child_text = fetch(loc)
                if child_text and ("<loc>" in child_text.lower()):
                    extract_urls(loc, collected, visited, log_lines)
                else:
                    collected.add(loc)
    else:
        for loc in locs:
            collected.add(loc)

    return collected


FALLBACK_PATHS = [
    "/sitemap.xml",
    "/sitemap_index.xml",
    "/sitemap/sitemap.xml",
    "/sitemap/index.xml",
    "/sitemaps/sitemap.xml",
    "/sitemap-index.xml",
]


def base_url(raw):
    """Return scheme://host for a raw domain or URL string."""
    u = raw.strip()
    if not u.startswith("http"):
        u = "https://" + u
    p = urlparse(u)
    return f"{p.scheme}://{p.netloc}"


@app.route("/extract", methods=["POST"])
def extract():
    data = request.get_json(force=True)
    raw = data.get("url", "").strip()
    override = data.get("override", "").strip()

    if not raw and not override:
        return jsonify({"error": "No URL provided"}), 400

    log_lines = []
    if override:
        sitemap_urls = [override]
    else:
        sitemap_urls = discover_sitemaps(raw, log_lines=log_lines)
    log_lines.append(f"Using {len(sitemap_urls)} sitemap(s)")

    collected = set()
    visited = set()
    for sitemap_url in sitemap_urls:
        extract_urls(sitemap_url, collected, visited, log_lines)

    return jsonify({
        "sitemap": ", ".join(sitemap_urls),
        "urls": sorted(collected),
        "count": len(collected),
        "sitemap_count": len(visited),
        "robots_sitemaps": len(sitemap_urls),
        "log": log_lines,
    })


@app.route("/robots")
def robots_txt():
    url = request.args.get("url", "").strip()
    if not url:
        return jsonify({"error": "No URL provided"}), 400
    if not url.startswith("http"):
        url = "https://" + url
    parsed = urlparse(url)
    base = f"{parsed.scheme}://{parsed.netloc}"
    text = fetch(base + "/robots.txt")
    if not text:
        return jsonify({"raw": "", "ok": False, "error": "Could not fetch robots.txt"})
    return jsonify({"raw": text, "ok": True})


@app.route("/ping")
def ping():
    return jsonify({"ok": True})


@app.route("/restart", methods=["POST"])
def restart():
    # Server is a daemon thread — no process restart needed.
    # Just return OK so the JS can reload the page and reset UI state.
    return jsonify({"ok": True})


def _fetch_text_file(url, timeout=20):
    """
    Fetch a plain-text file (e.g. llms.txt).
    Tries multiple user-agent profiles so WAFs / CDNs don't block it.
    Returns the decoded text or None.
    """
    header_sets = [
        HEADERS,
        HEADERS_GOOGLEBOT,
        HEADERS_CRAWLER,
        # Last resort: bare minimum headers
        {"User-Agent": "curl/8.4.0", "Accept": "*/*"},
    ]
    for hdrs in header_sets:
        try:
            r = requests.get(url, headers=hdrs, timeout=timeout, allow_redirects=True)
            if r.status_code == 404:
                return None  # definitely not there — no point retrying
            if not r.ok:
                continue  # try next header set
            text = _decode_response(r)
            if text and not text.lstrip().startswith("<"):
                return text  # valid plain-text content
            if text and text.lstrip().startswith("<"):
                return None  # HTML soft-404
        except requests.exceptions.SSLError:
            return None  # SSL errors won't be fixed by header changes
        except Exception:
            continue
    return None


@app.route("/ai-check")
def ai_check():
    url = request.args.get("url", "").strip()
    if not url:
        return jsonify({"error": "No URL provided"}), 400
    if not url.startswith("http"):
        url = "https://" + url
    parsed = urlparse(url)
    base = f"{parsed.scheme}://{parsed.netloc}"
    llms_text      = _fetch_text_file(base + "/llms.txt")
    llms_full_text = _fetch_text_file(base + "/llms-full.txt")

    return jsonify({
        "llms_txt": {
            "found": bool(llms_text),
            "content": llms_text[:3000] if llms_text else None,
        },
        "llms_full": {"found": bool(llms_full_text)},
    })


@app.route("/inspect-page")
def inspect_page():
    import json as _json
    import re as _re

    url = request.args.get("url", "").strip()
    if not url:
        return jsonify({"error": "No URL provided"}), 400
    if not url.startswith("http"):
        url = "https://" + url

    try:
        resp = requests.get(url, headers=HEADERS, timeout=20, allow_redirects=True)
        final_url   = resp.url
        status_code = resp.status_code
        html        = _decode_response(resp)

        # Cloudflare challenge — try bypass methods
        _was_cf_block = _is_cloudflare_block(resp)
        if _was_cf_block or not html:
            bypassed = _try_cloudflare_bypass(url, 25, None)
            if bypassed:
                html        = bypassed
                status_code = 200
                _was_cf_block = False
            elif not html:
                return jsonify({"error": f"Could not fetch page (HTTP {status_code})"}), 500
        # Final check: if the HTML we ended up with still looks like a CF challenge
        # (bypass returned challenge page), mark it as blocked.
        if not _was_cf_block and html:
            _cf_markers = ("just a moment", "enable javascript and cookies",
                           "checking your browser", "_cf_chl", "challenge-platform")
            _short_html = html[:3000].lower()
            if any(m in _short_html for m in _cf_markers) and len(html) < 20_000:
                _was_cf_block = True

        # ── Playwright fallback ───────────────────────────────────────────
        # If the page is CF-blocked or appears JS-rendered (body near-empty),
        # try the headless browser before falling back to the "no content" UI.
        _used_playwright = False
        _needs_pw = _was_cf_block or (html and len(html.split()) < 100)
        if _needs_pw and _PW_AVAILABLE:
            pw_html = _render_with_playwright(url)
            if pw_html and len(pw_html.split()) > 100:
                html          = pw_html
                final_url     = url
                status_code   = 200
                _was_cf_block = False
                _used_playwright = True

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    try:
        from bs4 import BeautifulSoup
        # lxml handles Shopify's inline <link>/<style> body injections correctly;
        # html.parser treats them as container elements, leaving body with only ~4
        # direct children and breaking _iter_blocks which walks body.children.
        try:
            import lxml as _lxml_check  # noqa
            _bs_parser = "lxml"
        except ImportError:
            _bs_parser = "html.parser"
        soup = BeautifulSoup(html, _bs_parser)

        title_tag = soup.find("title")
        title     = title_tag.get_text(strip=True) if title_tag else ""

        desc_tag  = soup.find("meta", attrs={"name": _re.compile(r"^description$", _re.I)})
        meta_desc = desc_tag.get("content", "").strip() if desc_tag else ""

        canon_tag = soup.find("link", rel=lambda r: r and "canonical" in r)
        canonical = canon_tag.get("href", "").strip() if canon_tag else ""

        robots_tag  = soup.find("meta", attrs={"name": _re.compile(r"^robots$", _re.I)})
        meta_robots = robots_tag.get("content", "").strip() if robots_tag else "index, follow (default)"

        og_title = (soup.find("meta", property="og:title") or {}).get("content", "")
        og_desc  = (soup.find("meta", property="og:description") or {}).get("content", "")
        og_img   = (soup.find("meta", property="og:image") or {}).get("content", "")

        # ── Images and links: count from FULL soup before any decomposition ──
        # Pre-filter removes product grids/noise which contain real images/links;
        # counting here ensures the totals reflect the full rendered page.
        _imgs_raw     = soup.find_all("img")
        images_total  = len(_imgs_raw)
        images_no_alt = sum(1 for i in _imgs_raw if not i.get("alt", "").strip())

        from urllib.parse import urljoin, urlparse as _up
        # Use final_url (post-redirect) so www./http: variants don't count as external
        _parsed_base = _up(final_url)
        _base_domain = _parsed_base.netloc
        internal_links, external_links = set(), set()
        for _a in soup.find_all("a", href=True):
            _href = _a["href"].strip()
            if _href.startswith(("#", "mailto:", "tel:")):
                continue
            _abs = urljoin(final_url, _href)
            _d   = _up(_abs).netloc
            if _d == _base_domain:
                internal_links.add(_abs)
            elif _d:
                external_links.add(_abs)

        body = soup.find("body")

        # ── Pre-filter: remove product grids / filter sidebars ───────────────
        # Uses simple substring matching on individual class tokens so BEM names
        # like "facets__form-wrapper" or "product-grid--template--xxx" are caught.

        # Substrings that mark filter/facet UI (check each class token individually)
        _FILTER_TOKENS = {'facet', 'filter', 'refinement', 'facets', 'filters'}

        # Substrings that mark product grid/listing containers
        _GRID_TOKENS   = {'product-grid', 'products-grid', 'product-list', 'product-loop',
                          'product-tile', 'product-card', 'collection-grid', 'collection-list',
                          'collection-loop', 'shop-grid', 'shop-loop', 'catalog-grid'}

        # Class tokens that explicitly mark editorial/SEO content — never strip
        _CONTENT_TOKENS = {'rte', 'rich-text', 'richtext', 'wysiwyg', 'editorial',
                           'seo-text', 'seo-content', 'cms-content', 'page-content'}

        # Class tokens that mark popup/modal/drawer UI overlays — always strip
        _POPUP_TOKENS = {'modal', 'popup', 'drawer', 'overlay', 'offcanvas',
                         'cart-notification', 'cookie-bar', 'cookie-banner', 'cookie-notice',
                         'announcement-bar', 'promo-bar', 'notification-bar',
                         'newsletter-popup', 'age-gate', 'lightbox',
                         # Blog/article listing widgets (not editorial content)
                         'blog-grid', 'blog-article', 'article-grid', 'recent-posts',
                         'latest-posts', 'post-grid', 'post-list',
                         # Footer navigation menus (divs with footer-menu/footer-nav classes)
                         'footer-menu', 'footer-nav', 'footer-links',
                         # Third-party review/testimonial widgets (Judge.me, Yotpo, Okendo etc.)
                         'jdgm', 'review-widget', 'reviews-widget', 'testimonial-widget',
                         'yotpo', 'okendo', 'stamped'}

        def _cls_tokens(el):
            return [c.lower() for c in (el.get('class') or [])]

        def _popup_match(token, popup_token):
            """Boundary-aware popup token check.

            Plain substring matching causes false positives on Shopify BEM names
            like 'product-modal__inner' (where 'modal' is part of a component
            name, not an actual popup).  A popup token only counts when it:
              - IS the full token  ("modal")
              - starts the token followed by '-' or '__'  ("modal-overlay", "modal__wrap")
              - ends the token preceded by '-'  ("search-modal", "cart-drawer")
            """
            if token == popup_token:
                return True
            if token.startswith(popup_token + '-') or token.startswith(popup_token + '__'):
                return True
            if token.endswith('-' + popup_token):
                return True
            return False

        def _is_noise_container(el):
            tokens = _cls_tokens(el)
            tokens_str = ' '.join(tokens)
            # <dialog> is always a UI overlay — strip unconditionally
            if el.name == 'dialog':
                return True
            # Protect elements explicitly marked as editorial content
            if any(ct in tokens for ct in _CONTENT_TOKENS):
                return False
            if any(ct in tokens_str for ct in _CONTENT_TOKENS):
                return False
            # Popup/modal/drawer/overlay — strip BEFORE the guardian check.
            # Uses boundary-aware matching to avoid false positives on Shopify BEM
            # names like 'product-modal__inner' or 'newsletter-popup__overlay'.
            if any(_popup_match(t, pt) for t in tokens for pt in _POPUP_TOKENS):
                return True
            # Protect anything that already contains long editorial paragraphs
            if any(len(p.get_text(strip=True)) > 100 for p in el.find_all('p', limit=3)):
                return False
            # Filter/facet UI: substring match per token (handles BEM __ separators)
            if any(ft in t for t in tokens for ft in _FILTER_TOKENS):
                return True
            # Product grid: require EXACT class token match OR element must be ul/ol.
            # Substring-only matching (e.g. "product-grid" in "product-grid--wrapper")
            # would decompose Shopify section wrappers that also contain editorial content.
            if any(gt in t for t in tokens for gt in _GRID_TOKENS):
                exact_match = any(gt == t for t in tokens for gt in _GRID_TOKENS)
                if exact_match or el.name in ('ul', 'ol'):
                    return True
            return False

        if body:
            _stripped_ids = set()
            for el in body.find_all(['div','ul','ol','section','aside','form','dialog']):
                if id(el) in _stripped_ids:
                    continue
                if _is_noise_container(el):
                    for desc in el.descendants:
                        _stripped_ids.add(id(desc))
                    el.decompose()

        # ── Headings: extracted AFTER pre-filter so noise headings are removed ──
        # (product grid H3s, blog article titles, footer nav headings etc. are gone)
        _heading_root = body if body else soup
        headings_ordered = [
            {"level": tag.name.upper(), "text": " ".join(tag.get_text(separator=" ").split())}
            for tag in _heading_root.find_all(["h1","h2","h3","h4","h5","h6"])
        ]
        headings = {
            "h1": [h["text"] for h in headings_ordered if h["level"] == "H1"],
            "h2": [h["text"] for h in headings_ordered if h["level"] == "H2"],
            "h3": [h["text"] for h in headings_ordered if h["level"] == "H3"],
            "ordered": headings_ordered,
        }

        # ── Content structure: run BEFORE stripping header/nav/footer ────
        def _clean(el):
            import re as _re2
            # Get text, collapse whitespace
            text = " ".join(el.get_text(separator=" ").split())
            # Fix BeautifulSoup space-before-punctuation artifact: "word ." → "word."
            text = _re2.sub(r'\s+([.,;:!?)\]»])', r'\1', text)
            text = _re2.sub(r'([(\[«])\s+', r'\1', text)
            return text

        def _safe_para(el):
            """Render a paragraph preserving <a href>, <strong>, <em>."""
            from bs4 import NavigableString as _NS2, Tag as _Tag3
            import re as _re5
            parts = []
            for child in el.children:
                if isinstance(child, _NS2):
                    parts.append(str(child))
                elif isinstance(child, _Tag3):
                    if child.name == "a":
                        href = (child.get("href") or "").strip()
                        txt  = child.get_text()
                        if href and txt:
                            parts.append(f'<a href="{href}" target="_blank" rel="noopener">{txt}</a>')
                        else:
                            parts.append(child.get_text())
                    elif child.name in ("strong","b"):
                        parts.append(f"<strong>{child.get_text()}</strong>")
                    elif child.name in ("em","i"):
                        parts.append(f"<em>{child.get_text()}</em>")
                    elif child.name == "br":
                        parts.append(" ")
                    else:
                        parts.append(child.get_text())
            raw = _re5.sub(r'\s+', ' ', "".join(parts)).strip()
            return _re5.sub(r'\s+([.,;:!?)\]»])', r'\1', raw)

        def _trunc(text, limit=800):
            """Truncate at sentence boundary so text never ends mid-word."""
            if len(text) <= limit:
                return text
            # Find last sentence-ending punctuation before limit
            cut = text.rfind('. ', 0, limit)
            if cut == -1:
                cut = text.rfind(' ', 0, limit)
            return text[:cut + 1].rstrip() + '…' if cut > 0 else text[:limit] + '…'

        # Noise token sets used by both _iter_blocks and Pass 2
        _IB_FILTER_TOKENS  = {'facet', 'filter', 'refinement'}
        _IB_CONTENT_TOKENS = {'rte', 'rich-text', 'richtext', 'wysiwyg',
                              'editorial', 'seo-', 'cms-content', 'page-content'}

        def _el_is_noise(el):
            """Return True if element looks like a filter/facet/popup UI container."""
            cls_list = [c.lower() for c in (el.get('class') or [])]
            if not cls_list:
                return False
            # Protect editorial content containers
            if any(ct in t for t in cls_list for ct in _IB_CONTENT_TOKENS):
                return False
            if any(ft in t for t in cls_list for ft in _IB_FILTER_TOKENS):
                return True
            return any(_popup_match(t, pt) for t in cls_list for pt in _POPUP_TOKENS)

        def _iter_blocks(el):
            """Yield block-level elements in document order, recursing into containers.
            header is treated as a container (not skipped) so H1s inside it are found.
            nav/footer/filter-sidebar are skipped to avoid noise.
            Leaf containers (divs with no block-level children but meaningful text)
            are yielded as paragraph substitutes for div-heavy modern layouts."""
            from bs4 import Tag as _Tag
            HEADING = {"h1","h2","h3","h4","h5","h6"}
            CONTENT = {"p","blockquote","pre"}
            LIST    = {"ul","ol"}
            SKIP    = {"script","style","nav","footer","dialog"}
            # header included as container so H1s inside <header> are not missed
            CONTAIN = {"div","section","article","main","aside","figure","form",
                       "header","details","summary","table","tbody","tr","td","th","li","dd"}
            BLOCK   = HEADING | CONTENT | LIST | CONTAIN
            for child in el.children:
                if not isinstance(child, _Tag):
                    continue
                if child.name in SKIP:
                    continue
                # Skip elements hidden from the page (hidden attr or aria-hidden)
                if child.get('hidden') is not None or child.get('aria-hidden') == 'true':
                    continue
                # Skip filter/facet UI containers at any nesting level
                if _el_is_noise(child):
                    continue
                if child.name in HEADING or child.name in CONTENT or child.name in LIST:
                    yield child
                elif child.name in CONTAIN:
                    # Check whether this container has any block-level children
                    has_block = any(
                        isinstance(c, _Tag) and c.name in BLOCK
                        for c in child.children
                    )
                    if has_block:
                        yield from _iter_blocks(child)
                    else:
                        # Leaf container — treat as paragraph if it has meaningful text
                        text = " ".join(child.get_text(separator=" ").split())
                        if text and len(text) > 20:
                            yield child
                else:
                    # Unknown/custom element (e.g. Shopify's <rte-formatter>,
                    # <product-card-link>, etc.) — recurse if it has block children,
                    # so editorial content inside custom elements is not silently dropped.
                    has_block = any(
                        isinstance(c, _Tag) and c.name in BLOCK
                        for c in child.children
                    )
                    if has_block:
                        yield from _iter_blocks(child)

        content_sections = []
        current = None
        root = body if body else soup
        for el in _iter_blocks(root):
            name = el.name
            if name in ("h1","h2","h3","h4","h5","h6"):
                text = _clean(el)
                if text:
                    current = {"level": int(name[1]), "heading": text, "content": []}
                    content_sections.append(current)
            elif name in ("ul","ol"):
                items = [_clean(li) for li in el.find_all("li", recursive=False)]
                items = [t for t in items if t][:12]
                if len(items) >= 2 and current is not None:
                    current["content"].append({"type": "list", "items": items})
            elif current is not None:
                # p, blockquote, pre, AND leaf-div containers all treated as content
                plain = _clean(el)
                if plain and len(plain) > 20:
                    current["content"].append({"type": "html", "html": _trunc(_safe_para(el), 1200), "text": _trunc(plain)})

        # ── Fallback: global scan to fill sections still empty after Pass 1 ──
        # Catches content in non-standard containers or partially JS-rendered pages.
        # Also scans leaf <div> elements for sites that don't use <p> tags.
        # Already-filled sections are not overwritten; seen_paras deduplicates.
        if content_sections:
            from bs4 import Tag as _TagFb
            heading_map = {s["heading"]: s for s in content_sections}
            # Seed seen_paras with content already captured so we don't duplicate
            seen_paras = {c["text"] for s in content_sections for c in s["content"] if c.get("type") in ("text","html")}
            current = None
            NOISE = {"nav","footer","script","style","noscript","dialog"}
            BLOCK_TAGS_SET = {"h1","h2","h3","h4","h5","h6","p","ul","ol","li",
                              "div","section","article","main","blockquote","pre"}
            for el in root.find_all(["h1","h2","h3","h4","h5","h6","p","div","ul","ol"]):
                name = el.name
                # Skip anything inside nav/footer/script to avoid noise
                if el.find_parent(NOISE):
                    continue
                if name in ("h1","h2","h3","h4","h5","h6"):
                    # Only update current if this heading is a known section heading
                    # — prevents nav/template headings from resetting current to None
                    matched = heading_map.get(_clean(el))
                    if matched is not None:
                        current = matched
                elif current is not None:
                    if name == "p":
                        plain = _clean(el)
                        if plain and len(plain) > 15 and plain not in seen_paras:
                            seen_paras.add(plain)
                            current["content"].append({"type": "html", "html": _trunc(_safe_para(el), 1200), "text": _trunc(plain)})
                    elif name == "div":
                        # Only process leaf divs (no block-level children) with meaningful text
                        has_block = any(
                            isinstance(c, _TagFb) and c.name in BLOCK_TAGS_SET
                            for c in el.children
                        )
                        if not has_block:
                            plain = _clean(el)
                            if plain and len(plain) > 15 and plain not in seen_paras:
                                seen_paras.add(plain)
                                current["content"].append({"type": "html", "html": _trunc(_safe_para(el), 1200), "text": _trunc(plain)})
                    elif name in ("ul","ol"):
                        # only top-level lists (not nested inside other lists)
                        if not el.find_parent(["ul","ol"]):
                            items = [_clean(li) for li in el.find_all("li")]
                            items = [t for t in items if t and len(t) > 5][:12]
                            if len(items) >= 2:
                                key = items[0]
                                if key not in seen_paras:
                                    seen_paras.add(key)
                                    current["content"].append({"type": "list", "items": items})

        # ── Pass 3: extract from __NEXT_DATA__ / embedded JSON payloads ────
        # Frameworks like Next.js store all page content as escaped HTML inside
        # a JSON script tag. Walk title+description and stack items, map by heading.
        import json as _json2
        def _extract_json_sections(soup_obj):
            results = {}  # heading_text -> [content blocks]
            def _text_from_html(html_str):
                return _clean(BeautifulSoup(html_str, "html.parser"))
            def _safe_inline(el):
                """Render element children preserving <a>, <strong>, <em> only."""
                from bs4 import NavigableString as _NS, Tag as _Tag2
                import re as _re4
                parts = []
                for child in el.children:
                    if isinstance(child, _NS):
                        parts.append(str(child))
                    elif isinstance(child, _Tag2):
                        if child.name == "a":
                            href = (child.get("href") or "").strip()
                            txt  = child.get_text()
                            if href and txt:
                                parts.append(f'<a href="{href}" target="_blank" rel="noopener">{txt}</a>')
                            else:
                                parts.append(child.get_text())
                        elif child.name in ("strong","b"):
                            parts.append(f"<strong>{child.get_text()}</strong>")
                        elif child.name in ("em","i"):
                            parts.append(f"<em>{child.get_text()}</em>")
                        elif child.name == "br":
                            parts.append(" ")
                        else:
                            parts.append(child.get_text())
                raw = "".join(parts)
                # collapse whitespace, fix space-before-punctuation
                raw = _re4.sub(r'\s+', ' ', raw).strip()
                raw = _re4.sub(r'\s+([.,;:!?)\]»])', r'\1', raw)
                return raw

            def _blocks_from_html(html_str):
                inner = BeautifulSoup(html_str, "html.parser")
                blocks = []
                for el in inner.find_all(["p","ul","ol","h1","h2","h3","h4","h5","h6"]):
                    if el.name == "p":
                        html_text = _safe_inline(el)
                        plain     = _clean(el)
                        if plain and len(plain) > 20:
                            blocks.append({"type": "html", "html": _trunc(html_text, 1200), "text": _trunc(plain)})
                    elif el.name in ("ul","ol") and not el.find_parent(["ul","ol"]):
                        items = []
                        for li in el.find_all("li"):
                            items.append({"html": _safe_inline(li), "text": _clean(li)})
                        items = [i for i in items if i["text"] and len(i["text"]) > 5][:12]
                        if items:
                            blocks.append({"type": "list", "items": items})
                # fallback: if no p/ul found, use full text
                if not blocks:
                    t = _clean(inner)
                    if t and len(t) > 20:
                        blocks.append({"type": "html", "html": _trunc(t, 1200), "text": _trunc(t)})
                return blocks

            def _str(val):
                """Safely coerce a JSON value to string; return '' for non-strings."""
                return val.strip() if isinstance(val, str) else ""

            def _walk(obj, depth=0):
                if depth > 20 or not obj: return
                if isinstance(obj, dict):
                    title = _str(obj.get("title") or obj.get("heading") or "")
                    # Strip inline HTML from title (e.g. "Healthy Eating, <br><br>Made Simple")
                    if title and "<" in title:
                        title = _clean(BeautifulSoup(title, "html.parser"))
                    # Try multiple content field names, only accept strings
                    desc = (_str(obj.get("description")) or _str(obj.get("body"))
                            or _str(obj.get("content")) or _str(obj.get("text")))
                    # Fallback: subtitle/cta for banner blocks
                    if not desc:
                        desc = " ".join(filter(None, [_str(obj.get("subtitle")), _str(obj.get("ctaText"))]))
                    if title and desc and len(desc) > 10:
                        results.setdefault(title, []).extend(_blocks_from_html(desc))
                    # stack: list of sub-items under a parent section
                    stack = obj.get("stack") or []
                    if isinstance(stack, list) and stack and title:
                        combined_blocks = []
                        for item in stack:
                            if isinstance(item, dict):
                                item_desc = _str(item.get("description"))
                                item_title = _str(item.get("title"))
                                if item_desc:
                                    if item_title:
                                        combined_blocks.append({"type": "text", "text": f"**{item_title}**"})
                                    combined_blocks.extend(_blocks_from_html(item_desc))
                        if combined_blocks:
                            results.setdefault(title, []).extend(combined_blocks)
                    for v in obj.values():
                        if isinstance(v, (dict, list)):
                            _walk(v, depth + 1)
                elif isinstance(obj, list):
                    for item in obj:
                        _walk(item, depth + 1)

            for script in soup_obj.find_all("script"):
                raw = script.string or ""
                if not raw or len(raw) < 100:
                    continue
                # __NEXT_DATA__ and similar JSON blobs
                try:
                    payload = _json2.loads(raw)
                    _walk(payload)
                except Exception:
                    pass
            return results

        if content_sections:
            json_data = _extract_json_sections(soup)
            heading_map_lower = {s["heading"].lower(): s for s in content_sections}

            def _norm_title(t):
                """Normalise for fuzzy matching: lowercase, strip site suffix, collapse spaces."""
                import re as _re3
                # Strip " | Brand", " - Brand", " — Brand" suffixes
                t = _re3.sub(r'\s*[\|\-—]\s*\S.*$', '', t)
                # Remove punctuation except spaces, lowercase
                t = _re3.sub(r'[^\w\s]', ' ', t.lower())
                # Sort words so "Best 7" matches "7 Best"
                return ' '.join(sorted(t.split()))

            norm_map = {_norm_title(s["heading"]): s for s in content_sections}
            seen_json = {c["text"] for s in content_sections for c in s["content"] if c.get("type") == "text"}
            for json_title, blocks in json_data.items():
                # Match by exact → lowercase → fuzzy normalised
                section = (heading_map.get(json_title)
                           or heading_map_lower.get(json_title.lower())
                           or norm_map.get(_norm_title(json_title)))
                if section is None:
                    continue
                for block in blocks:
                    if block.get("type") == "text":
                        t = block["text"]
                        if t not in seen_json:
                            seen_json.add(t)
                            section["content"].append(block)
                    else:
                        section["content"].append(block)

        # ── Pass 4: ordered-element walk for still-empty sections ──────────────
        # Collects all headings + paragraphs in flat document order (skipping nav/footer).
        # For each empty section at level N, gathers <p> elements that follow its heading
        # until the next heading at the SAME or HIGHER level (lower number).
        # This correctly handles two-column layouts where lower-level sidebar headings
        # (h4 "Client", h4 "Category") sit between an h1 and its content paragraphs.
        empty_sections = [s for s in content_sections if not s["content"]]
        if empty_sections:
            from bs4 import Tag as _TagP4
            NOISE_P4 = {"nav","footer","script","style","noscript","header","dialog"}

            # Flat list of all headings + paragraphs in document order, noise-filtered
            ordered_els = [
                el for el in root.find_all(["h1","h2","h3","h4","h5","h6","p","ul","ol"])
                if not el.find_parent(NOISE_P4)
            ]

            # Map heading text → its index in ordered_els (first occurrence)
            h_idx_map = {}
            for idx, el in enumerate(ordered_els):
                if el.name in {"h1","h2","h3","h4","h5","h6"}:
                    t = _clean(el)
                    if t and t not in h_idx_map:
                        h_idx_map[t] = idx

            p4_seen = {c["text"] for s in content_sections for c in s["content"] if "text" in c}

            for s in empty_sections:
                h_text  = s["heading"]
                h_level = s["level"]         # 1–6
                h_idx   = h_idx_map.get(h_text)
                if h_idx is None:
                    continue

                blocks = []
                for el in ordered_els[h_idx + 1:]:
                    name = el.name
                    if name in {"h1","h2","h3","h4","h5","h6"}:
                        el_level = int(name[1])
                        if el_level <= h_level:
                            break   # hit same/higher heading → stop
                        # lower-level heading (e.g. h4 sidebar label): skip but don't stop
                        continue
                    if name == "p":
                        t = _clean(el)
                        if t and len(t) > 15 and t not in p4_seen:
                            p4_seen.add(t)
                            blocks.append({"type": "html", "html": _trunc(_safe_para(el), 1200), "text": _trunc(t)})
                    elif name in ("ul","ol") and not el.find_parent(["ul","ol"]):
                        items = [_clean(li) for li in el.find_all("li")]
                        items = [i for i in items if i and len(i) > 5][:12]
                        if len(items) >= 2:
                            key = items[0]
                            if key not in p4_seen:
                                p4_seen.add(key)
                                blocks.append({"type": "list", "items": items})
                    if len(blocks) >= 10:
                        break

                s["content"].extend(blocks)

        # Strip noisy elements for word count (after content_sections is built)
        if body:
            for tag in body.find_all(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            body_text = body.get_text(separator=" ")
        else:
            body_text = soup.get_text(separator=" ")
        word_count = len([w for w in body_text.split() if w])

        # images_total, images_no_alt, internal_links, external_links
        # already computed above from the full pre-filter soup.

        SCHEMA_RULES = {
            "Article":             {"req": ["headline","author","datePublished"],       "rec": ["image","publisher","dateModified","description"]},
            "NewsArticle":         {"req": ["headline","author","datePublished"],       "rec": ["image","publisher"]},
            "BlogPosting":         {"req": ["headline","author","datePublished"],       "rec": ["image","publisher","dateModified"]},
            "Product":             {"req": ["name"],                                    "rec": ["description","image","offers","aggregateRating","brand"]},
            "FAQPage":             {"req": ["mainEntity"],                              "rec": []},
            "HowTo":               {"req": ["name","step"],                             "rec": ["description","image","totalTime"]},
            "Organization":        {"req": ["name"],                                    "rec": ["url","logo","sameAs","contactPoint"]},
            "LocalBusiness":       {"req": ["name","address"],                         "rec": ["telephone","openingHours","geo","url"]},
            "BreadcrumbList":      {"req": ["itemListElement"],                        "rec": []},
            "Event":               {"req": ["name","startDate"],                       "rec": ["endDate","location","description","organizer"]},
            "WebSite":             {"req": ["name"],                                    "rec": ["url","potentialAction"]},
            "WebPage":             {"req": ["name"],                                    "rec": ["url","description","breadcrumb"]},
            "Person":              {"req": ["name"],                                    "rec": ["jobTitle","url","sameAs"]},
            "SoftwareApplication": {"req": ["name","applicationCategory","operatingSystem"], "rec": ["offers","aggregateRating"]},
            "Recipe":              {"req": ["name","recipeIngredient","recipeInstructions"], "rec": ["image","author","totalTime","aggregateRating"]},
            "VideoObject":         {"req": ["name","description","thumbnailUrl","uploadDate"], "rec": ["duration","contentUrl","embedUrl"]},
            "JobPosting":          {"req": ["title","description","datePosted","hiringOrganization","jobLocation"], "rec": ["baseSalary","employmentType","validThrough"]},
            "Review":              {"req": ["reviewRating","author"],                  "rec": ["itemReviewed","reviewBody"]},
            "AggregateRating":     {"req": ["ratingValue","reviewCount"],              "rec": ["bestRating","worstRating"]},
            # E-commerce / store types
            "Store":               {"req": ["name","address"],                         "rec": ["url","telephone","openingHours","logo","sameAs"]},
            "OnlineStore":         {"req": ["name"],                                   "rec": ["url","logo","sameAs","address","telephone","description","acceptedPaymentMethod"]},
            "ItemList":            {"req": ["itemListElement"],                        "rec": ["name","description","numberOfItems"]},
            "ListItem":            {"req": ["position","item"],                        "rec": []},
            # Page types
            "CollectionPage":      {"req": ["name"],                                   "rec": ["url","description","mainEntity","breadcrumb"]},
            "AboutPage":           {"req": ["name"],                                   "rec": ["url","description","author"]},
            "ContactPage":         {"req": ["name"],                                   "rec": ["url","description"]},
            "SearchResultsPage":   {"req": ["name"],                                   "rec": ["url"]},
            "ImageObject":         {"req": ["contentUrl"],                             "rec": ["name","description","thumbnail","width","height"]},
            "Offer":               {"req": ["price","priceCurrency"],                  "rec": ["availability","priceValidUntil","url"]},
        }

        def validate_node(node):
            issues, warnings = [], []
            if not isinstance(node, dict):
                return {"type": "Unknown", "raw": str(node)[:500], "issues": ["Not a valid JSON-LD object"], "warnings": [], "valid": False, "score": 0, "known": False, "line_annotations": []}

            ctx_issue = False
            if not node.get("@context"):
                issues.append("Missing @context (should be 'https://schema.org')")
                ctx_issue = True
            elif "schema.org" not in str(node["@context"]):
                warnings.append(f"@context '{node['@context']}' may not be schema.org")
                ctx_issue = "warn"

            typ = node.get("@type", "")
            type_issue = not typ
            if type_issue:
                issues.append("Missing @type")

            type_str = (typ[0] if isinstance(typ, list) else str(typ)) if typ else "Unknown"
            rules = SCHEMA_RULES.get(type_str, {})
            req_set = set(rules.get("req", []))
            rec_set = set(rules.get("rec", []))
            for f in req_set:
                if f not in node:
                    issues.append(f"Missing required property: '{f}'")
            for f in rec_set:
                if f not in node:
                    warnings.append(f"Recommended property missing: '{f}'")

            if type_str == "FAQPage":
                for i, q in enumerate(node.get("mainEntity", []) or []):
                    if isinstance(q, dict):
                        if not q.get("name"):
                            issues.append(f"FAQ item {i+1}: missing 'name' (the question)")
                        aa = q.get("acceptedAnswer") or {}
                        if not aa.get("text"):
                            issues.append(f"FAQ item {i+1}: 'acceptedAnswer.text' missing")

            # ── Build line annotations ──────────────────────────────
            raw = _json.dumps(node, indent=2)
            lines = raw.split('\n')
            annotations = [None] * len(lines)
            prop_line = {}   # key -> first line index
            for i, line in enumerate(lines):
                m = _re.match(r'\s*"([^"]+)"\s*:', line)
                if m:
                    key = m.group(1)
                    if key not in prop_line:
                        prop_line[key] = i

            # @context line
            if "@context" in prop_line:
                annotations[prop_line["@context"]] = "error" if ctx_issue is True else ("warn" if ctx_issue == "warn" else "ok")
            # @type line
            if "@type" in prop_line:
                annotations[prop_line["@type"]] = "error" if type_issue else "ok"
            # required fields present → green; missing already in issues list (no line)
            for f in req_set:
                if f in prop_line:
                    annotations[prop_line[f]] = "ok"
            # recommended fields present → blue note
            for f in rec_set:
                if f in prop_line and annotations[prop_line[f]] is None:
                    annotations[prop_line[f]] = "rec"

            is_known = type_str in SCHEMA_RULES
            return {
                "type":             type_str,
                "raw":              raw[:4000],
                "issues":           issues,
                "warnings":         warnings,
                # Unknown types have no validation rules — mark valid=None so the
                # frontend can distinguish "not validated" from "validated & passing"
                "valid":            None if not is_known else len(issues) == 0,
                "score":            None if not is_known else max(0, 100 - len(issues)*20 - len(warnings)*5),
                "known":            is_known,
                "line_annotations": annotations,
            }

        schema_results = []
        for script in soup.find_all("script", type="application/ld+json"):
            raw_text = script.string or script.get_text()
            try:
                data = _json.loads(raw_text)
                nodes = data.get("@graph", None) if isinstance(data, dict) else None
                if nodes is not None:
                    ctx = data.get("@context", "")
                    for item in nodes:
                        if isinstance(item, dict):
                            item.setdefault("@context", ctx)
                        schema_results.append(validate_node(item))
                elif isinstance(data, list):
                    for item in data:
                        schema_results.append(validate_node(item))
                else:
                    schema_results.append(validate_node(data))
            except _json.JSONDecodeError as e:
                err_lines = raw_text.split('\n')
                ann = [None] * len(err_lines)
                if 1 <= e.lineno <= len(ann):
                    ann[e.lineno - 1] = "error"
                schema_results.append({"type": "Parse Error", "raw": raw_text[:2000], "issues": [f"Invalid JSON at line {e.lineno}, col {e.colno}: {e.msg}"], "warnings": [], "valid": False, "score": 0, "known": False, "line_annotations": ann})

        # ── Microdata detection ──────────────────────────────────────────────
        # Many Shopify themes use HTML microdata (itemscope/itemtype) instead of
        # JSON-LD. Extract top-level itemscope elements and report them.
        _seen_microdata_types = set()
        for el in soup.find_all(attrs={"itemscope": True}):
            # Skip nested itemscopes — only process top-level (no itemscope ancestor)
            if el.find_parent(attrs={"itemscope": True}):
                continue
            itemtype_url = (el.get("itemtype") or "").strip()
            # Normalise: "https://schema.org/Product" → "Product"
            md_type = itemtype_url.rsplit("/", 1)[-1].rsplit("#", 1)[-1] if itemtype_url else "Unknown"
            if md_type in _seen_microdata_types:
                continue
            _seen_microdata_types.add(md_type)
            # Collect itemprop names at any depth within this element
            props = {}
            for prop_el in el.find_all(attrs={"itemprop": True}):
                key = prop_el.get("itemprop", "").strip()
                if not key or key in props:
                    continue
                val = (prop_el.get("content") or prop_el.get("href") or
                       prop_el.get("src") or prop_el.get_text(strip=True))
                props[key] = str(val)[:200] if val else ""
            # Build a pseudo-JSON for display
            pseudo = {"@context": "https://schema.org", "@type": md_type}
            pseudo.update(props)
            result = validate_node(pseudo)
            result["format"] = "microdata"   # flag so UI can note it's not JSON-LD
            schema_results.append(result)

        def _detect_render_type(soup_obj, sections, wc):
            """Return render type string based on content extraction results."""
            total_blocks = sum(len(s["content"]) for s in sections) if sections else 0
            if sections and total_blocks >= 1:
                if total_blocks < len(sections) // 2:
                    return "partial_js"
                return "static"
            if wc >= 50:
                return "static"
            # No content — detect the JS framework to give a better hint
            raw_html = str(soup_obj)[:8000].lower()
            srcs = " ".join(s.get("src","") for s in soup_obj.find_all("script") if s.get("src")).lower()
            # Cloudflare Rocket Loader replaces script type with a hash prefix
            # e.g. type="81d772bab91062caecfad702-text/javascript"
            import re as _re_rl
            rocket_loader = any(
                _re_rl.match(r'[a-f0-9]{20,}-text/javascript', s.get("type",""))
                for s in soup_obj.find_all("script")
            )
            if rocket_loader:
                return "full_js_rocket"
            if "__next_data__" in raw_html or "/_next/" in srcs:
                return "full_js_nextjs"
            if "__nuxt" in raw_html or "/_nuxt/" in srcs:
                return "full_js_nuxt"
            if "ng-version" in raw_html or "/angular" in srcs:
                return "full_js_angular"
            if soup_obj.find(id="app") and ("react" in raw_html or "reactdom" in srcs):
                return "full_js_react"
            if soup_obj.find(id="app") or soup_obj.find(id="root"):
                return "full_js_spa"
            return "full_js"

        def _extract_page_meta(soup_obj):
            """Extract page metadata from GTM dataLayer pushes and meta tags."""
            import json as _jj, re as _rr
            meta = {}
            for script in soup_obj.find_all("script"):
                raw = script.string or ""
                for m in _rr.finditer(r'dataLayer\.push\s*\(\s*(\{.+?\})\s*\)', raw, _rr.DOTALL):
                    try:
                        d = _jj.loads(m.group(1))
                        if isinstance(d, dict):
                            if d.get("page_category") and not meta.get("page_category"):
                                meta["page_category"] = d["page_category"]
                            cpv = d.get("content_pageview", {})
                            if isinstance(cpv, dict):
                                if cpv.get("content_title") and not meta.get("content_title"):
                                    meta["content_title"] = cpv["content_title"]
                                if cpv.get("content_category") and not meta.get("content_category"):
                                    meta["content_category"] = cpv["content_category"]
                            # also check top-level event_category / page_type
                            for key in ("page_type", "event_category", "content_type"):
                                if d.get(key) and not meta.get(key):
                                    meta[key] = d[key]
                    except Exception:
                        pass
            return meta

        return jsonify({
            "url":            final_url,
            "status":         status_code,
            "title":          title,
            "meta_desc":      meta_desc,
            "canonical":      canonical,
            "meta_robots":    meta_robots,
            "og":             {"title": og_title, "description": og_desc, "image": og_img},
            "headings":       headings,
            "word_count":     word_count,
            "images_total":   images_total,
            "images_no_alt":  images_no_alt,
            "internal_links": len(internal_links),
            "external_links": len(external_links),
            "schema":           schema_results,
            "content_sections": content_sections,
            "js_rendered":      content_sections and sum(len(s["content"]) for s in content_sections) < len(content_sections) // 2,
            "render_type":      "cf_block" if _was_cf_block else _detect_render_type(soup, content_sections, word_count),
            "page_meta":        _extract_page_meta(soup),
            "used_playwright":  _used_playwright,
        })

    except Exception as e:
        return jsonify({"error": f"Parse error: {e}"}), 500


@app.route("/inspect-page-rendered")
def inspect_page_rendered():
    """Headless-browser render of a URL, returned as raw HTML for the
    frontend to hand back to /inspect-page via the override mechanism."""
    url = request.args.get("url", "").strip()
    if not url:
        return jsonify({"error": "No URL provided"}), 400
    if not url.startswith("http"):
        url = "https://" + url
    if not _PW_AVAILABLE:
        return jsonify({"error": "Playwright not installed"}), 503
    html = _render_with_playwright(url)
    if not html:
        return jsonify({"error": "Headless render failed"}), 500
    return jsonify({"html": html, "url": url})


@app.route("/export-geo", methods=["POST"])
def export_geo():
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        import io
    except ImportError:
        return jsonify({"error": "openpyxl not installed — run: pip install openpyxl"}), 500

    data        = request.get_json(force=True) or {}
    domain      = data.get("domain", "GEO-Audit")
    sections    = data.get("sections", [])
    entity_cats = data.get("entityStack", [])
    platform    = data.get("platform", {"headers": [], "rows": []})

    # ── Colour palette ──────────────────────────────────────────────
    C_NAV  = 'FF1A1F3D'  # deep navy
    C_MID  = 'FF1E2547'  # mid navy (section headers)
    C_WHT  = 'FFFFFFFF'
    C_SKY  = 'FF0EA5E9'  # sky blue (row numbers)
    C_DARK = 'FF1E293B'  # body text
    C_MUTE = 'FF94A3B8'  # muted subtitle
    C_CYAN = 'FF22D3EE'  # section score accent
    C_YLW  = 'FFFFF2CC'  # input yellow
    C_LITE = 'FFF0F4F8'  # alternating row bg
    C_RED  = 'FFEF4444'  # P1
    C_AMB  = 'FFF59E0B'  # P2
    C_GRN  = 'FF10B981'  # P3

    thin_side = Side(style='thin', color='FFCCCCCC')

    def mkfill(c):
        return PatternFill('solid', fgColor=c)

    def mkfont(bold=False, sz=9, color=C_DARK, italic=False):
        return Font(name='Arial', bold=bold, size=sz, color=color, italic=italic)

    def mkalign(h='center', wrap=True):
        return Alignment(horizontal=h, vertical='center', wrap_text=wrap)

    def mkborder():
        return Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)

    def cell_style(cell, value=None, bold=False, sz=9, fc=C_DARK, bg=None,
                   align_h='center', wrap=True, border=False, fmt=None, italic=False):
        if value is not None:
            cell.value = value
        cell.font      = mkfont(bold=bold, sz=sz, color=fc, italic=italic)
        cell.alignment = mkalign(h=align_h, wrap=wrap)
        if bg:
            cell.fill = mkfill(bg)
        if border:
            cell.border = mkborder()
        if fmt is not None:
            cell.number_format = fmt

    def prio_label(p):
        return 'P1 Critical' if p == 'P1' else ('P2 Important' if p == 'P2' else 'P3 Nice to Have')

    def prio_fc(label):
        return C_RED if 'P1' in label else (C_AMB if 'P2' in label else C_GRN)

    STATE_EMOJI = {'pass': '✅ Done', 'partial': '⚠️ Partial', 'fail': '❌ Missing'}

    wb = Workbook()

    # ═══════════════════════════════════════════════════════════════
    # Sheet 1 — GEO Page Audit
    # ═══════════════════════════════════════════════════════════════
    ws = wb.active
    ws.title = 'GEO Page Audit'
    ws.column_dimensions['A'].width = 4
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 44
    ws.column_dimensions['D'].width = 9
    ws.column_dimensions['E'].width = 17.29
    ws.column_dimensions['F'].width = 13
    ws.column_dimensions['G'].width = 9
    ws.column_dimensions['H'].width = 32

    def banner(ws, row, text, bg, fc=C_WHT, sz=14, bold=True, h=36):
        ws.merge_cells(f'A{row}:H{row}')
        ws.row_dimensions[row].height = h
        cell_style(ws[f'A{row}'], text, bold=bold, sz=sz, fc=fc, bg=bg)

    # Row 1 — Title
    banner(ws, 1, 'GEO PAGE AUDIT CHECKLIST  -  WEIGHTED SCORING', C_NAV)
    # Row 2 — Subtitle
    banner(ws, 2, 'Generative Engine Optimisation  |  Weighted scoring with signal detection',
           C_MID, fc=C_MUTE, sz=9, bold=False, h=21.75)

    # Row 3 — Client / URL
    ws.row_dimensions[3].height = 25.5
    ws.merge_cells('A3:B3')
    cell_style(ws['A3'], 'Client:', bold=True, sz=9, fc=C_WHT, bg=C_SKY)
    ws.merge_cells('C3:D3')
    cell_style(ws['C3'], domain, sz=9, fc=C_DARK, bg=C_WHT, border=True)
    cell_style(ws['E3'], 'URL:', bold=True, sz=9, fc=C_WHT, bg=C_SKY)
    ws.merge_cells('F3:H3')
    cell_style(ws['F3'], '', sz=9, fc=C_DARK, bg=C_WHT, border=True)

    # Row 4 — Focus Keyword / Date
    ws.row_dimensions[4].height = 25.5
    ws.merge_cells('A4:B4')
    cell_style(ws['A4'], 'Focus Keyword:', bold=True, sz=9, fc=C_WHT, bg=C_SKY)
    ws.merge_cells('C4:D4')
    cell_style(ws['C4'], '', sz=9, fc=C_DARK, bg=C_WHT, border=True)
    cell_style(ws['E4'], 'Date:', bold=True, sz=9, fc=C_WHT, bg=C_SKY)
    ws.merge_cells('F4:H4')
    cell_style(ws['F4'], '', sz=9, fc=C_DARK, bg=C_WHT, border=True)

    # Row 5 — Spacer
    ws.row_dimensions[5].height = 6
    for col in 'ABCDEFGH':
        ws[f'{col}5'].fill = mkfill(C_LITE)

    # Row 6 — Column headers
    ws.row_dimensions[6].height = 27.75
    col_heads = ['#', 'Check Item', 'Detail / What to Look For', 'Weight', 'Priority', 'Status', 'Score', 'Notes / Action']
    for i, h in enumerate(col_heads, 1):
        c = ws[f'{get_column_letter(i)}6']
        cell_style(c, h, bold=True, sz=10, fc=C_WHT, bg=C_NAV, border=True)

    # ── Write sections ──────────────────────────────────────────────
    cur = 7
    score_cells = []   # list of G{row} cell refs for overall formula
    sec_ranges  = []   # (item_start, item_end) per section

    total_all_weight = sum(
        item.get('weight', 1)
        for sec in sections for item in sec.get('items', [])
    )

    for sec in sections:
        items = sec.get('items', [])
        if not items:
            continue

        # Section weight % label
        sec_weight = sum(i.get('weight', 1) for i in items)
        pct = round(sec_weight / total_all_weight * 100) if total_all_weight else 0
        sec_title = f"{sec['title'].upper()}  (Weight: {pct}%)"

        # Section header
        ws.merge_cells(f'A{cur}:H{cur}')
        ws.row_dimensions[cur].height = 25.5
        cell_style(ws[f'A{cur}'], sec_title, bold=True, sz=11, fc=C_WHT, bg=C_MID)
        cur += 1

        item_start = cur
        for idx, item in enumerate(items):
            ws.row_dimensions[cur].height = 37.5
            row_bg = C_WHT if idx % 2 == 0 else C_LITE
            prio   = prio_label(item.get('priority', 'P2'))
            emoji  = STATE_EMOJI.get(item.get('state'), '')

            cell_style(ws[f'A{cur}'], idx + 1, bold=True, sz=9, fc=C_SKY,  bg=row_bg, border=True)
            cell_style(ws[f'B{cur}'], item.get('label', ''),   sz=9, fc=C_DARK, bg=row_bg, border=True)
            cell_style(ws[f'C{cur}'], item.get('detail', ''),  sz=9, fc=C_DARK, bg=row_bg, border=True)
            cell_style(ws[f'D{cur}'], item.get('weight', 1),   bold=True, sz=9, fc=C_DARK, bg=row_bg, border=True, fmt='0')
            cell_style(ws[f'E{cur}'], prio, bold=True, sz=9, fc=prio_fc(prio), bg=row_bg, border=True)
            cell_style(ws[f'F{cur}'], emoji, sz=9, fc=C_DARK, bg=C_YLW, border=True)
            # Score formula
            r = cur
            ws[f'G{r}'].value       = f'=IF(F{r}="✅ Done",D{r},IF(F{r}="⚠️ Partial",D{r}*0.5,IF(F{r}="❌ Missing",0,"")))'
            ws[f'G{r}'].font        = mkfont(bold=True, sz=9, color=C_DARK)
            ws[f'G{r}'].fill        = mkfill(row_bg)
            ws[f'G{r}'].alignment   = mkalign()
            ws[f'G{r}'].border      = mkborder()
            ws[f'G{r}'].number_format = '0'
            cell_style(ws[f'H{cur}'], '', sz=9, fc=C_DARK, bg=C_YLW, border=True)
            cur += 1

        item_end = cur - 1
        sec_ranges.append((item_start, item_end))

        # Section score footer
        ws.row_dimensions[cur].height = 27.75
        ws.merge_cells(f'A{cur}:E{cur}')
        cell_style(ws[f'A{cur}'], 'Section Score:', bold=True, sz=10, fc=C_CYAN, bg=C_NAV)
        cell_style(ws[f'F{cur}'], 'Earned:', sz=9, fc=C_MUTE, bg=C_NAV)
        sg = cur
        ws[f'G{sg}'].value          = (f'=SUMPRODUCT((F{item_start}:F{item_end}="✅ Done")*D{item_start}:D{item_end})'
                                       f'+SUMPRODUCT((F{item_start}:F{item_end}="⚠️ Partial")*D{item_start}:D{item_end})*0.5')
        ws[f'G{sg}'].font           = mkfont(bold=True, sz=12, color=C_CYAN)
        ws[f'G{sg}'].fill           = mkfill(C_NAV)
        ws[f'G{sg}'].alignment      = mkalign()
        ws[f'G{sg}'].number_format  = '0'
        ws[f'H{sg}'].value          = (f'=IF(COUNTA(F{item_start}:F{item_end})=0,"",ROUND(G{sg},1)'
                                       f'&" / "&ROUND(SUMPRODUCT((D{item_start}:D{item_end}>0)*D{item_start}:D{item_end}),0)&" pts")')
        ws[f'H{sg}'].font           = mkfont(sz=10, color=C_MUTE)
        ws[f'H{sg}'].fill           = mkfill(C_NAV)
        ws[f'H{sg}'].alignment      = mkalign()
        score_cells.append(f'G{sg}')
        cur += 1

        # Spacer
        ws.row_dimensions[cur].height = 6
        for col in 'ABCDEFGH':
            ws[f'{col}{cur}'].fill = mkfill(C_LITE)
        cur += 1

    # ── Overall GEO Score ───────────────────────────────────────────
    cur += 1  # blank separator
    ws.merge_cells(f'A{cur}:H{cur}')
    ws.row_dimensions[cur].height = 36
    cell_style(ws[f'A{cur}'], 'OVERALL GEO READINESS SCORE', bold=True, sz=14, fc=C_WHT, bg=C_NAV)
    cur += 1

    ws.row_dimensions[cur].height = 39.75
    ws.merge_cells(f'A{cur}:C{cur}')
    cell_style(ws[f'A{cur}'], 'Total Points Earned:', bold=True, sz=12, fc=C_WHT, bg=C_MID)

    total_formula = '+'.join(score_cells) if score_cells else '0'
    ws.merge_cells(f'D{cur}:E{cur}')
    ws[f'D{cur}'].value          = f'={total_formula}'
    ws[f'D{cur}'].font           = mkfont(bold=True, sz=18, color=C_SKY)
    ws[f'D{cur}'].fill           = mkfill(C_MID)
    ws[f'D{cur}'].alignment      = mkalign()
    ws[f'D{cur}'].number_format  = '0'

    d_ranges = '+'.join(f'SUM(D{s}:D{e})' for s, e in sec_ranges) if sec_ranges else '1'
    ws.merge_cells(f'F{cur}:H{cur}')
    ws[f'F{cur}'].value     = f'=IF(({total_formula})=0,"0%",TEXT(({total_formula})/({d_ranges}),"0%")&"  GEO Readiness")'
    ws[f'F{cur}'].font      = mkfont(bold=True, sz=12, color=C_CYAN)
    ws[f'F{cur}'].fill      = mkfill(C_MID)
    ws[f'F{cur}'].alignment = mkalign()
    cur += 1

    ws.merge_cells(f'A{cur}:H{cur}')
    ws.row_dimensions[cur].height = 27.75
    cell_style(ws[f'A{cur}'],
               'RATING:   90-100% = AI Citation Ready   |   70-89% = Strong Foundation   |   50-69% = Work Required   |   <50% = Major Gaps',
               sz=9, fc=C_CYAN, bg=C_NAV)
    cur += 1

    # P1 summary row
    p1_f_ranges = ','.join(f'F{s}:F{e}' for s, e in sec_ranges)
    p1_d_ranges = ','.join(f'D{s}:D{e}' for s, e in sec_ranges)
    ws.merge_cells(f'A{cur}:H{cur}')
    ws.row_dimensions[cur].height = 24
    if p1_f_ranges:
        ws[f'A{cur}'].value = (f'="P1 Critical items remaining: "&COUNTIFS({p1_f_ranges},"",'
                               f'{p1_d_ranges},">=1")&" of "&COUNTIF({p1_f_ranges},"<>")&" assessed"')
    else:
        ws[f'A{cur}'].value = 'P1 Critical items remaining: —'
    ws[f'A{cur}'].font      = mkfont(sz=9, color=C_MUTE)
    ws[f'A{cur}'].fill      = mkfill(C_NAV)
    ws[f'A{cur}'].alignment = mkalign()

    # ═══════════════════════════════════════════════════════════════
    # Sheet 2 — Entity Stacking
    # ═══════════════════════════════════════════════════════════════
    ws2 = wb.create_sheet('Entity Stacking')
    ws2.column_dimensions['A'].width = 4
    ws2.column_dimensions['B'].width = 28
    ws2.column_dimensions['C'].width = 50
    ws2.column_dimensions['D'].width = 12
    ws2.column_dimensions['E'].width = 32

    ws2.merge_cells('A1:E1')
    ws2.row_dimensions[1].height = 36
    cell_style(ws2['A1'], 'ENTITY STACKING CHECKLIST', bold=True, sz=14, fc=C_WHT, bg=C_NAV)
    ws2.merge_cells('A2:E2')
    ws2.row_dimensions[2].height = 21.75
    cell_style(ws2['A2'], f'Based on Floate framework  |  {domain}  |  Track entity presence across all platforms',
               sz=9, fc=C_MUTE, bg=C_MID)

    ws2.row_dimensions[3].height = 27.75
    for i, h in enumerate(['#', 'Entity Signal', 'What to Do', 'Status', 'URL / Notes'], 1):
        c = ws2[f'{get_column_letter(i)}3']
        cell_style(c, h, bold=True, sz=10, fc=C_WHT, bg=C_SKY, border=True)

    ecur = 4
    for cat in entity_cats:
        ws2.merge_cells(f'A{ecur}:E{ecur}')
        ws2.row_dimensions[ecur].height = 25.5
        cell_style(ws2[f'A{ecur}'], cat.get('title', '').upper(), bold=True, sz=11, fc=C_WHT, bg=C_MID)
        ecur += 1
        for idx, item in enumerate(cat.get('items', [])):
            ws2.row_dimensions[ecur].height = 31.5
            row_bg  = C_WHT if idx % 2 == 0 else C_LITE
            checked = item.get('checked', False)
            cell_style(ws2[f'A{ecur}'], idx + 1,           bold=True, sz=9, fc=C_SKY,  bg=row_bg, border=True)
            cell_style(ws2[f'B{ecur}'], item.get('label',''),  sz=9, fc=C_DARK, bg=row_bg, border=True)
            cell_style(ws2[f'C{ecur}'], item.get('note',''),   sz=9, fc=C_DARK, bg=row_bg, border=True)
            cell_style(ws2[f'D{ecur}'], '✅ Live' if checked else '', sz=9, fc=C_DARK, bg=C_YLW, border=True)
            cell_style(ws2[f'E{ecur}'], '', sz=9, fc=C_DARK, bg=C_YLW, border=True)
            ecur += 1

    ws2.row_dimensions[ecur].height = 15.75
    ecur += 1
    ws2.merge_cells(f'A{ecur}:E{ecur}')
    ws2.row_dimensions[ecur].height = 27.75
    ws2[f'A{ecur}'].value     = f'=COUNTIF(D4:D{ecur-1},"✅ Live")&" / "&COUNTA(B4:B{ecur-1})&" live"'
    ws2[f'A{ecur}'].font      = mkfont(sz=10, color=C_CYAN)
    ws2[f'A{ecur}'].fill      = mkfill(C_NAV)
    ws2[f'A{ecur}'].alignment = mkalign()

    # ═══════════════════════════════════════════════════════════════
    # Sheet 3 — Platform Citation Matrix
    # ═══════════════════════════════════════════════════════════════
    ws3 = wb.create_sheet('Platform Citation Matrix')
    pm_headers = platform.get('headers', [])
    pm_rows    = platform.get('rows', [])

    ws3.column_dimensions['A'].width = 24
    for i in range(len(pm_headers)):
        ws3.column_dimensions[get_column_letter(i + 2)].width = 22

    ws3.merge_cells(f'A1:{get_column_letter(len(pm_headers) + 1)}1')
    ws3.row_dimensions[1].height = 36
    cell_style(ws3['A1'], 'PLATFORM CITATION MATRIX', bold=True, sz=14, fc=C_WHT, bg=C_NAV)
    ws3.merge_cells(f'A2:{get_column_letter(len(pm_headers) + 1)}2')
    ws3.row_dimensions[2].height = 21.75
    cell_style(ws3['A2'], 'How each AI platform retrieves, weights, and cites content  |  Use to prioritise optimisation effort',
               sz=9, fc=C_MUTE, bg=C_MID)

    ws3.row_dimensions[3].height = 27.75
    cell_style(ws3['A3'], 'Signal', bold=True, sz=10, fc=C_WHT, bg=C_NAV, border=True)
    for i, h in enumerate(pm_headers, 2):
        c = ws3[f'{get_column_letter(i)}3']
        cell_style(c, h, bold=True, sz=10, fc=C_WHT, bg=C_NAV, border=True)

    for ridx, row_data in enumerate(pm_rows):
        r = ridx + 4
        ws3.row_dimensions[r].height = 31.5
        row_bg = C_WHT if ridx % 2 == 0 else C_LITE
        cell_style(ws3[f'A{r}'], row_data.get('signal', ''), bold=True, sz=9, fc=C_DARK, bg=row_bg, border=True)
        for ci, val in enumerate(row_data.get('vals', []), 2):
            cell_style(ws3[f'{get_column_letter(ci)}{r}'], val, sz=9, fc=C_DARK, bg=row_bg, border=True)

    # ── Serialise and return ────────────────────────────────────────
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f'{domain}-GEO-Audit.xlsx'
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


def _docx_add_hyperlink(para, url, text):
    """Insert a clickable hyperlink run into a python-docx paragraph."""
    from docx.oxml.ns import qn
    import docx.oxml as oxml
    r_id = para.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = oxml.OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    wr = oxml.OxmlElement("w:r")
    rPr = oxml.OxmlElement("w:rPr")
    rStyle = oxml.OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    wr.append(rPr)
    t = oxml.OxmlElement("w:t")
    t.text = text
    wr.append(t)
    hl.append(wr)
    para._p.append(hl)


def _docx_fill_html(para, html_str):
    """Parse HTML and add runs/hyperlinks to an existing paragraph."""
    from html.parser import HTMLParser
    from docx.shared import RGBColor
    import html as html_mod

    class _P(HTMLParser):
        def __init__(self):
            super().__init__()
            self._href = None
            self._buf  = []
            self._bold = False
            self._ital = False

        def _commit(self):
            t = "".join(self._buf).replace("\xa0", " ")
            self._buf = []
            return t

        def _flush(self):
            t = self._commit()
            if t:
                r = para.add_run(t)
                r.bold   = self._bold
                r.italic = self._ital

        def _flush_link(self):
            t = self._commit()
            if not t:
                return
            if self._href:
                try:
                    _docx_add_hyperlink(para, self._href, t)
                    return
                except Exception:
                    pass
            r = para.add_run(t)
            r.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)

        def handle_starttag(self, tag, attrs):
            d = dict(attrs)
            if tag == "a":
                self._flush(); self._href = d.get("href", "")
            elif tag in ("b", "strong"):
                self._flush(); self._bold = True
            elif tag in ("i", "em"):
                self._flush(); self._ital = True

        def handle_endtag(self, tag):
            if tag == "a":
                self._flush_link(); self._href = None
            elif tag in ("b", "strong"):
                self._flush(); self._bold = False
            elif tag in ("i", "em"):
                self._flush(); self._ital = False

        def handle_data(self, data):
            self._buf.append(data)

        def handle_entityref(self, name):
            self._buf.append(html_mod.unescape(f"&{name};"))

        def handle_charref(self, name):
            c = chr(int(name[1:], 16) if name.startswith("x") else int(name))
            self._buf.append(c)

    p = _P()
    p.feed(html_str)
    p._flush()


@app.route("/save-bulk-files", methods=["POST"])
def save_bulk_files():
    import subprocess
    import json as _json
    import traceback

    # Parse request body safely
    try:
        data = request.get_json(force=True, silent=True) or {}
    except Exception:
        data = {}

    files  = data.get("files", [])
    folder = os.path.expanduser(data.get("folder") or "~/Desktop/CrawlSync Exports")

    try:
        os.makedirs(folder, exist_ok=True)
    except Exception as e:
        return jsonify({"saved": 0, "errors": [f"Cannot create folder: {e}"], "folder": folder})

    # Import python-docx
    try:
        import docx as _docx
        from docx import Document
        from docx.shared import RGBColor
        template = os.path.join(os.path.dirname(_docx.__file__), "templates", "default.docx")
        if not os.path.exists(template):
            template = None
    except Exception as imp_err:
        return jsonify({"saved": 0, "errors": [f"ImportError: {imp_err}"], "folder": folder})

    saved  = []
    errors = []
    for f in files:
        try:
            path = os.path.join(folder, f.get("name", "page.docx"))
            doc  = Document(template) if template else Document()

            # Source URL as hyperlink
            url = f.get("url", "") or ""
            if url:
                p = doc.add_paragraph()
                try:
                    _docx_add_hyperlink(p, url, url)
                except Exception:
                    r = p.add_run(url)
                    r.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)
                doc.add_paragraph()

            if f.get("error"):
                doc.add_paragraph(f'Error: {f["error"]}')
            else:
                sections = _json.loads(f.get("sections_json", "[]") or "[]")
                seen_h = set()
                for s in sections:
                    hk = (s.get("heading", "") or "").strip().lower()
                    if hk and hk in seen_h:
                        continue
                    if hk:
                        seen_h.add(hk)
                    level   = max(1, min(int(s.get("level", 1) or 1), 9))
                    heading = (s.get("heading", "") or "").strip()
                    if heading:
                        doc.add_heading(heading, level=level)
                    for c in (s.get("content", []) or []):
                        ctype = c.get("type", "")
                        if ctype == "text":
                            txt = (c.get("text", "") or "").strip()
                            if txt:
                                doc.add_paragraph(txt)
                        elif ctype == "html":
                            html_src = (c.get("html", "") or "").strip()
                            if html_src:
                                p = doc.add_paragraph()
                                _docx_fill_html(p, html_src)
                        elif ctype == "list":
                            for item in (c.get("items", []) or []):
                                ih = item.get("html", "") if isinstance(item, dict) else ""
                                it = item.get("text", "") if isinstance(item, dict) else str(item)
                                try:
                                    lp = doc.add_paragraph(style="List Bullet")
                                except Exception:
                                    lp = doc.add_paragraph()
                                if ih:
                                    _docx_fill_html(lp, ih)
                                elif it:
                                    lp.add_run(it.strip())

            doc.save(path)
            saved.append(path)
        except Exception:
            errors.append(traceback.format_exc()[-300:])

    if saved:
        import platform as _platform
        try:
            if _platform.system() == "Windows":
                os.startfile(folder)
            elif _platform.system() == "Darwin":
                subprocess.Popen(["open", folder])
            else:
                subprocess.Popen(["xdg-open", folder])
        except Exception:
            pass
    return jsonify({"saved": len(saved), "errors": errors, "folder": folder})


@app.route("/pick-folder")
def pick_folder_dialog():
    """Windows folder-picker via tkinter (used when running in Edge app mode, no pywebview)."""
    import platform as _platform
    if _platform.system() != "Windows":
        return jsonify({"error": "pick-folder endpoint is Windows-only"}), 400
    try:
        import tkinter as _tk
        from tkinter import filedialog as _fd
        root = _tk.Tk()
        root.withdraw()
        root.wm_attributes("-topmost", True)
        folder = _fd.askdirectory(title="Choose folder to save files")
        root.destroy()
        return jsonify({"folder": folder or None})
    except Exception as e:
        return jsonify({"error": str(e), "folder": None})


@app.route("/quit", methods=["POST"])
def quit_app():
    threading.Thread(target=lambda: (time.sleep(0.3), os._exit(0))).start()
    return jsonify({"ok": True})


if __name__ == "__main__":
    print("CrawlSync server running → http://localhost:5000")
    app.run(port=5000, debug=False)
