#!/usr/bin/env python3
"""
CrawlSync launcher.
Starts the local extraction server and opens the tool in a native WebView window.

macOS build:   ./build_app.sh   → dist/CrawlSync.app  + CrawlSync.dmg
Windows build: build_app.bat    → dist/CrawlSync/CrawlSync.exe  + CrawlSync-Windows.zip
"""
import os
import sys
import threading
import time
import urllib.request
import importlib
import importlib.util
import webview

# Force-load sitemap_server from the data file (not the frozen module)
# so that the bundled .py is always up to date.
def _load_sitemap_server():
    # Debug log — written to home dir (always writable, cross-platform).
    # Wrapped in try/except so a missing/unwritable path never crashes the app.
    def _log(msg):
        try:
            log = os.path.join(os.path.expanduser("~"), "_crawlsync_debug.txt")
            with open(log, "a", encoding="utf-8") as f:
                f.write(msg + "\n")
        except Exception:
            pass

    _log(f"_MEIPASS: {getattr(sys, '_MEIPASS', 'NOT SET')}")
    _log(f"__file__: {__file__}")

    candidates = []
    if hasattr(sys, "_MEIPASS"):
        mp = sys._MEIPASS
        candidates += [
            mp,
            os.path.normpath(os.path.join(mp, "..", "Frameworks")),
            os.path.normpath(os.path.join(mp, "..", "Resources")),
        ]
    candidates.append(os.path.dirname(os.path.abspath(__file__)))

    for base in candidates:
        fp = os.path.join(base, "sitemap_server.py")
        exists = os.path.exists(fp)
        _log(f"  {fp} -> exists={exists}")
        if exists:
            try:
                spec = importlib.util.spec_from_file_location("sitemap_server", fp)
                mod  = importlib.util.module_from_spec(spec)
                sys.modules["sitemap_server"] = mod
                spec.loader.exec_module(mod)
                routes = [str(r) for r in mod.app.url_map.iter_rules()]
                _log(f"  LOADED OK. Routes: {routes}")
                return mod
            except Exception as e:
                import traceback
                _log(f"  LOAD FAILED: {e}\n{traceback.format_exc()}")

    _log("Falling back to frozen import")
    import sitemap_server
    return sitemap_server

sitemap_server = _load_sitemap_server()


class CrawlSyncAPI:
    def pick_folder(self):
        """Show a native folder picker and return the chosen path, or None."""
        result = webview.windows[0].create_file_dialog(
            webview.FOLDER_DIALOG,
            directory=os.path.expanduser('~'),
        )
        return result[0] if result else None

    @staticmethod
    def _add_hyperlink(para, url, text):
        """Insert a clickable hyperlink run into para."""
        from docx.oxml.ns import qn
        import docx.oxml as oxml
        r_id = para.part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True,
        )
        hl = oxml.OxmlElement('w:hyperlink')
        hl.set(qn('r:id'), r_id)
        wr = oxml.OxmlElement('w:r')
        rPr = oxml.OxmlElement('w:rPr')
        rStyle = oxml.OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        wr.append(rPr)
        t = oxml.OxmlElement('w:t')
        t.text = text
        wr.append(t)
        hl.append(wr)
        para._p.append(hl)

    @staticmethod
    def _fill_para_from_html(para, html_str):
        """Parse an HTML snippet and add runs/hyperlinks to an existing paragraph."""
        from html.parser import HTMLParser
        from docx.shared import RGBColor
        import html as html_mod

        class _P(HTMLParser):
            def __init__(self):
                super().__init__()
                self._href = None
                self._buf  = []
                self._bold = False
                self._ital = False

            def _commit(self):
                text = ''.join(self._buf).replace('\xa0', ' ')
                self._buf = []
                return text

            def _flush(self):
                text = self._commit()
                if text:
                    run = para.add_run(text)
                    run.bold   = self._bold
                    run.italic = self._ital

            def _flush_link(self):
                text = self._commit()
                if not text:
                    return
                if self._href:
                    try:
                        CrawlSyncAPI._add_hyperlink(para, self._href, text)
                        return
                    except Exception:
                        pass
                run = para.add_run(text)
                run.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)

            def handle_starttag(self, tag, attrs):
                d = dict(attrs)
                if tag == 'a':
                    self._flush()
                    self._href = d.get('href', '')
                elif tag in ('b', 'strong'):
                    self._flush(); self._bold = True
                elif tag in ('i', 'em'):
                    self._flush(); self._ital = True

            def handle_endtag(self, tag):
                if tag == 'a':
                    self._flush_link(); self._href = None
                elif tag in ('b', 'strong'):
                    self._flush(); self._bold = False
                elif tag in ('i', 'em'):
                    self._flush(); self._ital = False

            def handle_data(self, data):
                self._buf.append(data)

            def handle_entityref(self, name):
                self._buf.append(html_mod.unescape(f'&{name};'))

            def handle_charref(self, name):
                c = chr(int(name[1:], 16) if name.startswith('x') else int(name))
                self._buf.append(c)

        p = _P()
        p.feed(html_str)
        p._flush()

    def save_files(self, folder, files):
        """Write files as .docx into folder, open in Finder, return count saved."""
        import subprocess
        import traceback
        from docx.shared import RGBColor
        try:
            import docx as _docx
            from docx import Document
            template = os.path.join(os.path.dirname(_docx.__file__), 'templates', 'default.docx')
            if not os.path.exists(template):
                template = None
        except Exception:
            template = None
            Document = None

        saved = 0
        log   = []
        for f in files:
            try:
                name = f.get('name', 'page.docx') if isinstance(f, dict) else 'page.docx'
                path = os.path.join(folder, name)
                if Document:
                    doc = Document(template) if template else Document()

                    # Source URL as hyperlink
                    url = (f.get('url', '') if isinstance(f, dict) else '') or ''
                    if url:
                        p = doc.add_paragraph()
                        try:
                            self._add_hyperlink(p, url, url)
                        except Exception:
                            run = p.add_run(url)
                            run.font.color.rgb = RGBColor(0x00, 0x56, 0xD2)
                        doc.add_paragraph()

                    # Sections — passed as JSON string to avoid pywebview bridge issues
                    import json as _json
                    raw_json = (f.get('sections_json', '[]') if isinstance(f, dict) else '[]') or '[]'
                    sections = _json.loads(raw_json)

                    seen = set()
                    for s in sections:
                        key = (s.get('heading', '') or '').strip().lower()
                        if key and key in seen:
                            continue
                        if key:
                            seen.add(key)

                        level   = max(1, min(int(s.get('level', 1) or 1), 9))
                        heading = (s.get('heading', '') or '').strip()
                        if heading:
                            doc.add_heading(heading, level=level)

                        for c in (s.get('content', []) or []):
                            ctype = c.get('type', '')
                            if ctype == 'text':
                                txt = (c.get('text', '') or '').strip()
                                if txt:
                                    doc.add_paragraph(txt)
                            elif ctype == 'html':
                                html_src = (c.get('html', '') or '').strip()
                                if html_src:
                                    p = doc.add_paragraph()
                                    self._fill_para_from_html(p, html_src)
                            elif ctype == 'list':
                                for item in (c.get('items', []) or []):
                                    if isinstance(item, dict):
                                        item_html = item.get('html', '')
                                        item_txt  = item.get('text', '')
                                    else:
                                        item_html = ''
                                        item_txt  = str(item)
                                    try:
                                        lp = doc.add_paragraph(style='List Bullet')
                                    except Exception:
                                        lp = doc.add_paragraph()
                                    if item_html:
                                        self._fill_para_from_html(lp, item_html)
                                    elif item_txt:
                                        lp.add_run(item_txt.strip())
                    doc.save(path)
                else:
                    # Fallback: plain text
                    with open(path.replace('.docx', '.txt'), 'w', encoding='utf-8') as fh:
                        fh.write(f.get('content', '') if isinstance(f, dict) else '')
                saved += 1
            except Exception:
                log.append(traceback.format_exc())

        if saved:
            import platform as _platform
            try:
                if _platform.system() == 'Windows':
                    os.startfile(folder)
                elif _platform.system() == 'Darwin':
                    subprocess.Popen(['open', folder])
                else:
                    subprocess.Popen(['xdg-open', folder])
            except Exception:
                pass
        if log:
            with open(os.path.join(folder, '_errors.txt'), 'w') as fh:
                fh.write('\n\n'.join(log))
        return saved


def run_server():
    sitemap_server.app.run(host="127.0.0.1", port=5050, debug=False, use_reloader=False)


def wait_for_server(timeout=10):
    """Poll until Flask is ready or timeout is reached."""
    deadline = time.time() + timeout
    while time.time() < deadline:
        try:
            urllib.request.urlopen("http://127.0.0.1:5050/ping", timeout=0.5)
            return True
        except Exception:
            time.sleep(0.1)
    return False


def main():
    threading.Thread(target=run_server, daemon=True).start()

    if not wait_for_server():
        # Show a basic error page if server never came up
        webview.create_window("CrawlSync", html="<h2>Server failed to start.</h2>", width=400, height=200)
        webview.start()
        return

    api = CrawlSyncAPI()
    webview.create_window(
        "CrawlSync",
        "http://127.0.0.1:5050/",
        width=1400,
        height=900,
        min_size=(800, 600),
        js_api=api,
    )
    webview.start()
    # Window closed — daemon thread exits with the process


if __name__ == "__main__":
    main()
